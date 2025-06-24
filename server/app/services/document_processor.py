# server/app/services/document_processor.py
import os
import tempfile
import pdfplumber
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from .embedder import get_embeddings
from .pinecone_client import upsert_documents
from .s3_client import upload_document_to_s3
from ..database import SessionLocal
from ..models import Document
import mimetypes
from sqlalchemy.orm import Session


def extract_text_from_pdf(file_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file_bytes)
        tmp.flush()
        with pdfplumber.open(tmp.name) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        os.remove(tmp.name)
        return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes)
        tmp.flush()
        doc = docx.Document(tmp.name)
        text = "\n".join(para.text for para in doc.paragraphs)
        os.remove(tmp.name)
        return text.strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")


def process_and_store_document(user_id: str, filename: str, file_bytes: bytes, db: Session):
    # 1. Extract text
    text = extract_text(file_bytes, filename)
    if not text:
        raise ValueError("No extractable text found.")

    # 2. Chunk text
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = splitter.split_text(text)

    # 3. Embed text
    embeddings = get_embeddings(chunks)

    # 4. Upload embeddings to Pinecone
    metadata = {"filename": filename, "user_id": user_id}
    pinecone_namespace = upsert_documents(str(user_id), chunks, embeddings, metadata)

    # 5. Upload to S3
    content_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
    file_path = upload_document_to_s3(user_id, filename, file_bytes, content_type)

    # 6. Save metadata to DB
    db = SessionLocal()
    document = Document(
        user_id=user_id,
        filename=filename,
        file_path=file_path,
        file_size=len(file_bytes),
        content_type=content_type,
        pinecone_namespace=pinecone_namespace
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    db.close()

    return {
        "filename": filename,
        "num_chunks": len(chunks),
        "s3_key": file_path,
        "document_id": document.id
    }
